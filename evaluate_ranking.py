#!/usr/bin/python3

import json
import seaborn as sns
from matplotlib import cm
from matplotlib.colors import ListedColormap, LinearSegmentedColormap
import matplotlib.colors as colors
from scipy.stats import spearmanr
import pylab
import scipy.cluster.hierarchy as sch


from scipy.stats import pearsonr, friedmanchisquare
from scikit_posthocs import posthoc_nemenyi_friedman
from mlflow.tracking.client import MlflowClient
from mlflow.entities import ViewType


from docx import Document
from docx.shared import Inches

from datetime import datetime
import pandas as pd
import random
import numpy as np
import time
#import pycm
import shutil
import pathlib
import os
import math
import sys
import random
from matplotlib import pyplot
import matplotlib.pyplot as plt
import time
import copy
import random
import pickle
import tempfile
import itertools
import multiprocessing
import socket
from glob import glob
from collections import OrderedDict
import logging
import mlflow
from typing import Dict, Any
import hashlib
import json

from pprint import pprint
from sklearn.metrics import confusion_matrix
from sklearn.metrics import roc_curve, auc, roc_auc_score, precision_recall_curve
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, classification_report

from loadData import *
from utils import *
from parameters import *
from evaluate_utils import *



### parameters
TrackingPath = "/data/results/radFS/mlrun.benchmark"
nCV = 10



def plot_Ranking_Matrix (dList, results):
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'
    document.add_heading('Supplemental 4')
    document.add_paragraph(' ')
    document.add_heading('Ranking', level = 2)
    document.add_paragraph(' ')

    fSels = sorted(list(set(results["FSel"].values)))
    nList = sorted(list(set(results["nFeatures"].values)))

    def getRanking (d, fselA, nList):
        fv = []
        for N in sorted(nList, reverse  = True):
            subdata = results.query("Dataset == @d and FSel == @fselA and nFeatures == @N and Clf == 'Constant'")
            z = []
            for m in range(nCV):
                apath = os.path.join(subdata.iloc[0]["Path"], "FPattern_" + str(m) + ".json")
                with open(apath) as f:
                    expData = json.load(f)
                z.append(list(expData.values()))
            z = np.asarray(z)
            fv.append(z)
        fv = np.mean(np.concatenate(fv), axis = 0)
        return (fv)

    fSelList = fSels
    cMatList = {}
    for d in dList:
        if os.path.exists("./results/rank_" + d + ".npy") == False:
            rMat = np.ones( (len(fSelList), len(fSelList) ) )
            rMat = pd.DataFrame(rMat, index = fSelList, columns = fSelList)
            for fselA in fSelList:
                if fselA == "None":
                    continue
                for fselB in fSelList:
                    if fselB == "None":
                        continue
                    if fselA == fselB:
                        continue
                    #spearmanr( getRanking(d, fselA, nList), getRanking(d, fselB, nList))[0]
                    rMat.at[fselA, fselB] = spearmanr(getRanking(d, fselA, nList), getRanking(d, fselB, nList))[0]
            cMat  = rMat.drop(["None"], axis = 1)
            cMat  = cMat.drop(["None"], axis = 0)
            pickle.dump (cMat, open("./results/rank_" + d + ".npy","wb"))
        else:
            print ("Restoring rank results")
            cMat = pickle.load(open("./results/rank_" + d + ".npy", "rb"))
        cMatList[d] = cMat
    fullMat = np.mean(np.dstack(list(cMatList.values())), axis = 2)
    zMat = cMat.copy()
    zMat[:] = fullMat



    def plotDendro (document, zMat, d = None):
        D = zMat.copy()
        fig = pylab.figure(figsize=(20,20), dpi = DPI)

        # Compute and plot second dendrogram.
        ax2 = fig.add_axes([0.3,0.71,0.6,0.2])
        Y = sch.linkage(cMat, method='single')
        Z2 = sch.dendrogram(Y)
        ax2.set_xticks([])
        ax2.set_yticks([])

        # Plot distance matrix.
        axmatrix = fig.add_axes([0.3,0.1,0.6,0.6])
        idx1 = Z2['leaves']
        D = D.iloc[idx1]
        D = D[D.keys()[idx1]]
        im = axmatrix.matshow(D, norm=MidpointNormalize(midpoint=0,vmin=-1.0, vmax=1.0),\
                    aspect='auto', origin='lower', cmap=pylab.cm.RdBu)
        axmatrix.set_yticks([])
        axmatrix.xaxis.tick_bottom()
        xaxis = np.arange(len(D.keys()))
        axmatrix.set_xticks(xaxis)
        axmatrix.set_yticks(xaxis)
        axmatrix.set_xticklabels(D.keys(), rotation = 90,   fontsize = 22)
        axmatrix.set_yticklabels(D.keys(), ha = "right", fontsize = 22)

        # Plot colorbar.
        axcolor = fig.add_axes([0.91,0.1,0.02,0.6])
        cbar = pylab.colorbar(im, cax=axcolor)
        cbar.ax.tick_params(labelsize=22)
        plt.tight_layout()
        #fig.show()
        if d is None:
            paragraph = document.add_paragraph('All Datasets')
            fig.savefig('results/Ranking.png', bbox_inches='tight')
            # this plot is not there after revision
            #fig.savefig("./paper/Figure_4.png", facecolor = 'w', bbox_inches='tight')
            document.add_picture("./results/Ranking.png", width=Inches(6.0))
        else:
            paragraph = document.add_paragraph('Dataset ' + d)
            fig.savefig('results/Ranking_' + d + '.png', bbox_inches='tight')
            document.add_picture("./results/Ranking_"+d+".png", width=Inches(6.0))
        document.add_page_break()
        pass


    # dendros
    plotDendro (document, zMat)
    for d in dList:
        plotDendro (document, cMatList[d], d)
    document.save('./paper/Supplemental_5.docx')
    plt.close('all')
    pass






if __name__ == "__main__":
    print ("Hi.")

    # load data first
    mlflow.set_tracking_uri(TrackingPath)

    # datasets
    dList = [ "Carvalho2018", "Hosny2018A", "Hosny2018B", "Hosny2018C", "Ramella2018",   "Toivonen2019",
        "Keek2020", "Li2020", "Park2020", "Song2020" , ]

    # obtain results
    results = getResults (dList)

     # ranking
    plot_Ranking_Matrix (dList, results)


#
