#!/usr/bin/python3

import json
import seaborn as sns
from matplotlib import cm
from matplotlib.colors import ListedColormap, LinearSegmentedColormap
import matplotlib.colors as colors
from scipy.stats import spearmanr
import pylab
import scipy.cluster.hierarchy as sch

import cv2
from PIL import ImageDraw, ImageFont
from PIL import Image

from scipy.stats import pearsonr, friedmanchisquare
from scikit_posthocs import posthoc_nemenyi_friedman
from mlflow.tracking.client import MlflowClient
from mlflow.entities import ViewType


from docx import Document
from docx.shared import Inches

from datetime import datetime
import pandas as pd
import random
import numpy as np
import time
#import pycm
import shutil
import pathlib
import os
import math
import sys
import random
from matplotlib import pyplot
import matplotlib.pyplot as plt
import time
import copy
import random
import pickle
import tempfile
import itertools
import multiprocessing
import socket
from glob import glob
from collections import OrderedDict
import logging
import mlflow
from typing import Dict, Any
import hashlib
import json

from pprint import pprint
from sklearn.metrics import confusion_matrix
from sklearn.metrics import roc_curve, auc, roc_auc_score, precision_recall_curve
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, classification_report

from loadData import *
from utils import *
from parameters import *
from evaluate_utils import *



### parameters
TrackingPath = "/data/results/radFS/mlrun.benchmark"
nCV = 10




def addText (finalImage, text = '', org = (0,0), fontFace = '', fontSize = 12, color = (255,255,255)):
     tmpImg = finalImage
     pil_im = Image.fromarray(tmpImg)
     draw = ImageDraw.Draw(pil_im)
     font = ImageFont.truetype(fontFace + ".ttf", fontSize)
     draw.text(org, text, font=font, fill = color)
     tmpImg = np.array(pil_im)
     return (tmpImg.copy())



if __name__ == "__main__":
    print ("Hi.")

    # load data first
    mlflow.set_tracking_uri(TrackingPath)

    # datasets
    dList = [ "Carvalho2018", "Hosny2018A", "Hosny2018B", "Hosny2018C", "Ramella2018",   "Toivonen2019",
        "Keek2020", "Li2020", "Park2020", "Song2020" , ]

    # obtain results
    results = getResults (dList)

    # just read both figures and merge them
    def join_plots():
        fontFace = "Arial"

        imB = cv2.imread("./paper/Figure_2B.png")
        imB = np.hstack([255*np.ones(( imB.shape[0],int(imB.shape[0]*0.03), 3), dtype = np.uint8),imB])
        imB = addText (imB, "(b)", (0,0), fontFace, 18, color=(0,0,0))
        imB = np.hstack([255*np.ones(( imB.shape[0],int(imB.shape[0]*0.05), 3), dtype = np.uint8),imB])

        imA = cv2.imread("./paper/Figure_1C.png")
        imA = np.hstack([255*np.ones(( imA.shape[0],int(imA.shape[0]*0.03), 3), dtype = np.uint8),imA])
        imA = addText (imA, "(a)", (0,0), fontFace, 18, color= (0,0,0))

        img = np.hstack([imA, imB])
        #Image.fromarray(img)
        cv2.imwrite("./paper/Figure_8.png", img)
    join_plots()

    # create supplemental
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'
    document.add_heading('Supplemental 8')
    document.add_paragraph(' ')
    document.add_heading('Complexity Trade-offs', level = 2)
    document.add_page_break()
    document.add_heading('Training times vs AUC', level = 2)
    document.add_paragraph(' ')

    for d in dList + ["all"]:
        paragraph = document.add_paragraph('Dataset ' + d)
        img = cv2.imread("./results/Training_Times_vs_AUC_"+d+".png")
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        document.add_picture("./results/Training_Times_vs_AUC_"+d+".png", width=Inches(6.0))
        document.add_paragraph()

    document.add_page_break()
    document.add_heading('Stability vs AUC', level = 3)
    document.add_paragraph(' ')

    for d in dList + ["all"]:
        paragraph = document.add_paragraph('Dataset ' + d)
        img = cv2.imread("./results/Stability_vs_AUC_"+d+".png")
        img = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
        document.add_picture("./results/Stability_vs_AUC_"+d+".png", width=Inches(6.0))
        document.add_page_break()

    document.save('./paper/Supplemental_8.docx')

#
