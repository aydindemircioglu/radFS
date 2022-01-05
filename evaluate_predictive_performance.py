#!/usr/bin/python3

import json
import seaborn as sns
from matplotlib import cm
from matplotlib.colors import ListedColormap, LinearSegmentedColormap
import matplotlib.colors as colors
from scipy.stats import spearmanr
import pylab
import scipy.cluster.hierarchy as sch


from scipy.stats import pearsonr, friedmanchisquare
from scikit_posthocs import posthoc_nemenyi_friedman
from mlflow.tracking.client import MlflowClient
from mlflow.entities import ViewType


from docx import Document
from docx.shared import Inches

from datetime import datetime
import pandas as pd
import random
import numpy as np
import time
#import pycm
import shutil
import pathlib
import os
import math
import sys
import random
from matplotlib import pyplot
import matplotlib.pyplot as plt
import time
import copy
import random
import pickle
import tempfile
import itertools
import multiprocessing
import socket
from glob import glob
from collections import OrderedDict
import logging
import mlflow
from typing import Dict, Any
import hashlib
import json

from pprint import pprint
from sklearn.metrics import confusion_matrix
from sklearn.metrics import roc_curve, auc, roc_auc_score, precision_recall_curve
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, classification_report

from loadData import *
from utils import *
from parameters import *
from evaluate_utils import *


### parameters
TrackingPath = "/data/results/radFS/mlrun.benchmark"
nCV = 10





# die besten 10 configs fuer jedes FSel als boxplot pro dataset
def plot_Overall_Performance (dList, results):
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'
    document.add_heading('Supplemental 7')
    document.add_paragraph(' ')
    document.add_heading('Predictive Performance', level = 2)
    document.add_paragraph(' ')

    cTable = []
    fTables = {}
    for d in dList:
        fTable = []
        fSels = sorted(list(set(results["FSel"].values)))
        nList = sorted(list(set(results["nFeatures"].values)))
        aTable = results.query("Dataset == @d")
        aTable = aTable.sort_values("AUC_mean", ascending = False).reset_index(drop = True).copy()
        lmax = aTable.iloc[0]["AUC_mean"]
        aTable["AUC_score"] = lmax - aTable["AUC_mean"]

        for fsel in fSels:
            dTable = aTable.query("Dataset == @d and FSel == @fsel")
            if fsel == "None":
                dTable = aTable.query("nFeatures == 1")
            #fTable.append(dTable.iloc[0:10])
            fTable.append(dTable.iloc[0:int(dTable.shape[0]*0.05)])


        fTable = pd.concat(fTable)
        fTables[d] = fTable

        sOrder = fTable.groupby (["FSel"])["AUC_score"].median()
        sOrder = sOrder.sort_values(ascending = False)

        tmp = fTable.copy()
        cTable.append(tmp)

    # over all datasets
    cTable = pd.concat(cTable)
    fTables["all"] =  cTable.copy()

    for d in dList + ["all"]:
        paragraph = document.add_paragraph('Dataset ' + d)

        plotData = fTables[d].copy()

        sOrder = plotData.groupby (["FSel"])["AUC_score"].median()
        sOrder = sOrder.sort_values(ascending = False)

        fig, ax = plt.subplots(figsize = (20,10), dpi = DPI)
        sns.set(style='white')
        sns.boxplot(x = 'FSel', y = 'AUC_score', data= plotData,   order = sOrder.keys())#, whis = 50.0)
        sns.despine()

        ax.set_xticklabels(sOrder.keys(), rotation = 45, ha = "right", fontsize = 22)
        ax.yaxis.set_tick_params ( labelsize= 22)
        ax.set_xlabel ("Feature Selection Method", fontsize = 26)
        ax.set_ylabel ("Mean Relative AUC", fontsize = 26)
        #plt.legend(bbox_to_anchor=(1.05, 1), loc=2, borderaxespad=0.)
        ax.invert_xaxis()

        plt.tight_layout()
        fig.savefig("./results/Predictive_Performance_"+d+".png", facecolor = 'w')
        if d == "all":
            fig.savefig("./paper/Figure_6.png", facecolor = 'w')
        document.add_picture("./results/Predictive_Performance_"+d+".png", width=Inches(6.0))
        document.add_page_break()

    document.save('./paper/Supplemental_7.docx')
    plt.close('all')
    pass



# die besten 10 configs fuer jedes FSel als boxplot pro dataset
def plot_Outperforming (dList, results):
    fSels = sorted(list(set(results["FSel"].values)))
    rMat = np.zeros( (len(fSels), len(fSels) ) )
    rMat = pd.DataFrame(rMat, index = fSels, columns = fSels)
    for d in dList:
        dTable = results.query("Dataset == @d")
        dTable = dTable.sort_values("AUC_mean", ascending = False)[["AUC_mean", "FSel"]].reset_index(drop = True)
        dTable["rank"] = dTable["AUC_mean"].rank(method = "average", ascending = False)

        for idxA in range(len(fSels)):
            for idxB in range(len(fSels)):
                if idxA < idxB:
                    fselA = fSels[idxA]
                    fselB = fSels[idxB]
                    rankA = dTable.query("FSel == @fselA").iloc[0]["rank"]
                    rankB = dTable.query("FSel == @fselB").iloc[0]["rank"]
                    if rankA < rankB:
                        rMat.at[fselA, fselB] = rMat.at[fselA, fselB] + 1
                    if rankA > rankB:
                        rMat.at[fselB, fselA] = rMat.at[fselB, fselA] + 1
                    if rankA == rankB:
                        rMat.at[fselA, fselB] = rMat.at[fselA, fselB] + 0.5
                        rMat.at[fselB, fselA] = rMat.at[fselB, fselA] + 0.5
    # sort martix by rowsums
    rMat["sums"] = rMat.sum(axis = 1)
    rMat = rMat.sort_values(by = "sums", ascending = False)
    rMat = rMat[rMat.index] # removes sums again

    # remove identity
    for fselA in fSels:
        rMat.at[fselA, fselA] = None

    fig, ax = plt.subplots(figsize = (23,14), dpi = DPI)
    sns.set(style='white')
    ax = sns.heatmap(rMat, annot = True, cmap = pylab.cm.PiYG, annot_kws={"fontsize":21})
    cbar = ax.collections[0].colorbar
    cbar.ax.tick_params(labelsize=22)
    plt.tight_layout()

    ax.invert_xaxis()
    ax.set_xticklabels(rMat.keys(), rotation = 45, ha = "right", fontsize = 22)
    ax.set_yticklabels(rMat.keys(), rotation = 0, ha = "right", fontsize = 22)
    ax.yaxis.set_tick_params ( labelsize= 22)
    ax.set_xlabel ("Losses", fontsize = 22)
    ax.set_ylabel ("Wins", fontsize = 22)
    fig.savefig("./paper/Figure_7.png", facecolor = 'w', bbox_inches='tight')

    plt.close('all')
    pass



def plot_Top_Models (dList, results):
    fSels = sorted(list(set(results["FSel"].values)))
    clfs = sorted(list(set(results["Clf"].values)))
    clfs = sorted([k for k in clfs if k != "Constant"])
    rMat = np.zeros( (len(clfs), len(fSels) ) )
    rMat = pd.DataFrame('', index = clfs, columns = fSels)
    vMat = pd.DataFrame(0.0, index = clfs, columns = fSels)
    for d in dList:
        for f in fSels:
            for c in clfs:
                dTable = results.query("Dataset == @d and FSel == @f and Clf == @c")
                dTable = dTable.sort_values("AUC_mean", ascending = False).reset_index(drop = True)
                rMat.at[c, f] = str(round(dTable.iloc[0]["AUC_mean"],2)) + "\n(+/- " + str(round(dTable.iloc[0]["AUC_std"],2)) + ")"
                vMat.at[c, f] = dTable.iloc[0]["AUC_mean"]
        if 1 == 1:
            fig, ax = plt.subplots(figsize = (40,10), dpi = 50)
            sns.set(style='white')
            vMat
            sns.heatmap(vMat, annot = rMat,  cmap = "Reds", fmt = '', annot_kws={"fontsize":21}, cbar = True)
            ax.set_xticklabels(vMat.keys(), rotation = 45, ha = "right", fontsize = 21)
            ax.set_yticklabels(vMat.index, rotation = 0, ha = "right", fontsize = 21)
            ax.yaxis.set_tick_params ( labelsize= 21)
            ax.set_xlabel ("", fontsize = 19)
            ax.set_ylabel ("", fontsize = 19)
            ax.set_title("", fontsize = 24)
            plt.tight_layout()
            fig.savefig("./results/Predictive_Performance_Best_Combination_"+d+".png", facecolor = 'w')

    plt.close('all')
    pass






if __name__ == "__main__":
    print ("Hi.")

    # load data first
    mlflow.set_tracking_uri(TrackingPath)

    # datasets
    dList = [ "Carvalho2018", "Hosny2018A", "Hosny2018B", "Hosny2018C", "Ramella2018",   "Toivonen2019",
        "Keek2020", "Li2020", "Park2020", "Song2020" , ]

    # obtain results
    results = getResults (dList)

    # overall performance
    plot_Overall_Performance (dList, results)
    plot_Outperforming (dList, results)

#
