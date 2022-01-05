#!/usr/bin/python3

import json
import seaborn as sns
from matplotlib import cm
from matplotlib.colors import ListedColormap, LinearSegmentedColormap
import matplotlib.colors as colors
from scipy.stats import spearmanr
import pylab
import scipy.cluster.hierarchy as sch


from scipy.stats import pearsonr, friedmanchisquare
from scikit_posthocs import posthoc_nemenyi_friedman
from mlflow.tracking.client import MlflowClient
from mlflow.entities import ViewType


from docx import Document
from docx.shared import Inches

from datetime import datetime
import pandas as pd
import random
import numpy as np
import time
#import pycm
import shutil
import pathlib
import os
import math
import sys
import random
from matplotlib import pyplot
import matplotlib.pyplot as plt
import time
import copy
import random
import pickle
import tempfile
import itertools
import multiprocessing
import socket
from glob import glob
from collections import OrderedDict
import logging
import mlflow
from typing import Dict, Any
import hashlib
import json

from pprint import pprint
from sklearn.metrics import confusion_matrix
from sklearn.metrics import roc_curve, auc, roc_auc_score, precision_recall_curve
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, classification_report

from loadData import *
from utils import *
from parameters import *
from evaluate_utils import *



### parameters
TrackingPath = "/data/results/radFS/mlrun.benchmark"
nCV = 10



def plot_Timings(results, dList, closeDoc = True):
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'
    document.add_heading('Supplemental 2')
    document.add_paragraph(' ')
    document.add_heading('Training times', level = 2)
    document.add_paragraph(' ')

    allResults = results.copy()
    if 1 == 1:
        for d in dList + ["all"]:
            if d == "all":
                plotData = allResults.copy()
            else:
                plotData = allResults.query("Dataset == @d").copy()

            paragraph = document.add_paragraph('Dataset ' + d)

            fSels = sorted(list(set(plotData["FSel"].values)))
            nList = sorted(list(set(plotData["nFeatures"].values)))

            selTimes = []
            for fsel in fSels:
                for n in nList:
                    tmpData = plotData.query("FSel == @fsel and nFeatures == @n").copy()
                    allTimes = list(tmpData[[f for f in tmpData.keys() if "Fsel_Time" in f]].values.flat)
                    for t in allTimes:
                        selTimes.append({"FSel": fsel, "N": n, "Time": t })
            selTimes = pd.DataFrame(selTimes)
            selTimes["N"] = selTimes["N"].astype("category")
            selTimes = selTimes.groupby(["FSel"]).mean()
            selTimes = selTimes.sort_values("Time")
            selTimes = selTimes.reset_index()
            if 1 == 1:
                greysBig = cm.get_cmap('Greys', 512)
                newcmp = ListedColormap(greysBig(np.linspace(0.35, 0.9, 256)))

                fig, ax = plt.subplots(figsize = (20,10), dpi = DPI)
                sns.set(style='white')

                sns.barplot (x = 'FSel', y = 'Time',  data = selTimes)
                sns.despine()

                ax.set_yscale('log')
                ax.xaxis.set_tick_params ( labelsize= 22)
                plt.xticks(rotation=45, ha='right')
                ax.yaxis.set_tick_params ( labelsize= 22)
                ax.set_xlabel ("Feature Selection Method", fontsize = 26)
                ax.set_ylabel ("Mean Training Time [s]", fontsize = 26)
                plt.tight_layout()
                fig.savefig("./results/Training_Times_Mean_"+d+".png", facecolor = 'w')
                if d == "all":
                    fig.savefig("./paper/Figure_1.png", facecolor = 'w')
                document.add_picture("./results/Training_Times_Mean_"+d+".png", width=Inches(6.0))
                document.add_page_break()

    if closeDoc == True:
        document.save('./paper/Supplemental_2.docx')
    plt.close('all')
    return document



def plot_Timing_Combinations(results, dList, document = None, closeDoc = False):
    if document is None:
        document = Document()
        font = document.styles['Normal'].font
        font.name = 'Arial'
        document.add_heading('Supplemental 2')
        document.add_paragraph(' ')

    document.add_heading('Training times (Combinations)', level = 2)
    document.add_paragraph(' ')

    if 1 == 1:
        for d in ["all"] + dList:
            if d == "all":
                plotData = results.copy()
            else:
                plotData = results.query("Dataset == @d").copy()

            paragraph = document.add_paragraph('Dataset ' + d)


            Clfs = sorted(list(set(plotData["Clf"].values)))
            selTimes = []
            for clf in Clfs:
                tmpData = plotData.query("Clf == @clf").copy()
                tmpData["FSel_Time_Total"] = tmpData[[f for f in tmpData.keys() if "Fsel_Time" in f]].sum(axis=1)
                tmpData["Clf_Time_Total"] = tmpData[[f for f in tmpData.keys() if "Clf_Time" in f]].sum(axis=1)
                tmpData["Time_Total"] = tmpData["Clf_Time_Total"]+tmpData["FSel_Time_Total"]
                tmpData["FSel_Time_Rel"] = tmpData["FSel_Time_Total"]/tmpData["Time_Total"]
                tmpData["Clf_Time_Rel"] = tmpData["Clf_Time_Total"]/tmpData["Time_Total"]
                tmpData
                tmpData = tmpData[["FSel_Time_Rel", "Clf_Time_Rel", "FSel_Time_Total", "Clf_Time_Total"]].mean(axis = 0)
                tmpData["Clf"] = clf
                selTimes.append(tmpData)
            selTimes = pd.DataFrame(selTimes)
            selTimes



            fSels = sorted(list(set(plotData["FSel"].values)))
            selTimes = []
            for fsel in fSels:
                tmpData = plotData.query("FSel == @fsel").copy()
                tmpData["FSel_Time_Total"] = tmpData[[f for f in tmpData.keys() if "Fsel_Time" in f]].sum(axis=1)
                tmpData["Clf_Time_Total"] = tmpData[[f for f in tmpData.keys() if "Clf_Time" in f]].sum(axis=1)
                tmpData["Time_Total"] = tmpData["Clf_Time_Total"]+tmpData["FSel_Time_Total"]
                tmpData["FSel_Time_Rel"] = tmpData["FSel_Time_Total"]/tmpData["Time_Total"]
                tmpData["Clf_Time_Rel"] = tmpData["Clf_Time_Total"]/tmpData["Time_Total"]
                tmpData
                tmpData = tmpData[["FSel_Time_Rel", "Clf_Time_Rel", "FSel_Time_Total", "Clf_Time_Total"]].mean(axis = 0)
                tmpData["FSel"] = fsel
                selTimes.append(tmpData)
            selTimes = pd.DataFrame(selTimes)
            selTimes


            # time combinations
            rMat = np.zeros( (len(fSels), len(Clfs) ) )
            rMat = pd.DataFrame(rMat, index = fSels, columns = Clfs)
            for clf in Clfs:
                for fsel in fSels:
                    tmpData = plotData.query("Clf == @clf and FSel == @fsel").copy()
                    tmpData["FSel_Time_Total"] = tmpData[[f for f in tmpData.keys() if "Fsel_Time" in f]].sum(axis=1)
                    tmpData["Clf_Time_Total"] = tmpData[[f for f in tmpData.keys() if "Clf_Time" in f]].sum(axis=1)
                    tmpData["Time_Total"] = tmpData["Clf_Time_Total"]+tmpData["FSel_Time_Total"]
                    rMat.at[fsel, clf] = tmpData["Time_Total"].mean()

            # just for lda
            max2 = sorted(list(rMat.values.flat))[-2]
            labels = rMat.copy()
            idx = rMat > 3*max2
            labels = np.array(labels)
            labels = labels.round(2).astype(str)

            if np.sum(idx.values) > 0:
                rMat[rMat > max2] = 3*max2
                fixed = labels[idx][0] + "*"
                labels[idx] = fixed
                labels = labels.transpose()
                rMat = rMat.transpose()

            if 1 == 1:
                fig, ax = plt.subplots(figsize = (20,15), dpi = DPI)
                sns.set(style='white')
                sns.heatmap(rMat, annot = labels,  cmap = "Reds", fmt = '', annot_kws={"fontsize":21}, cbar = False)
                ax.set_xticklabels(rMat.keys(), rotation = 45, ha = "right", fontsize = 21)
                ax.set_yticklabels(rMat.index, rotation = 0, ha = "right", fontsize = 21)
                ax.yaxis.set_tick_params ( labelsize= 21)
                ax.set_xlabel ("", fontsize = 19)
                ax.set_ylabel ("", fontsize = 19)
                ax.set_title("", fontsize = 24)
                plt.tight_layout()
                fig.savefig("./results/Training_Times_Combinations_"+d+".png", facecolor = 'w')
                if d == "all":
                    fig.savefig("./paper/Figure_2.png", facecolor = 'w')
                document.add_picture("./results/Training_Times_Combinations_"+d+".png", width=Inches(6.0))
                document.add_page_break()

    if closeDoc == True:
        document.save('./paper/Supplemental_2.docx')
    plt.close('all')
    return document



def plot_Timing_AUC(results, dList):

    fSels = sorted(list(set(results["FSel"].values)))
    nList = sorted(list(set(results["nFeatures"].values)))

    nTable = []
    for d in dList:
        subdata = results.query("Dataset == @d")
        subdata = subdata.sort_values("AUC_mean", ascending = False).iloc[0:1].reset_index(drop = True).copy()
        lmax = subdata.iloc[0]["AUC_mean"]
        subdata = results.query("Dataset == @d").copy()
        subdata["AUC_score"] = lmax - subdata["AUC_mean"]
        nTable.append(subdata)
    mainTable = pd.concat(nTable)
    mainTable["FTime"] = mainTable[[k for k in mainTable.keys() if "Fsel_Ti" in k]].sum(axis = 1)
    applyLOG = True
    if applyLOG == True:
        mainTable["FTime"] = np.log(1 + mainTable["FTime"])

    mainTable = mainTable[["FTime", "AUC_score", "FSel", "Dataset"]].copy()
    for d in dList + ["all"]:
        if d == "all":
            mTable = mainTable.groupby(["FSel", "Dataset"]).mean()
        else:
            nTable = mainTable.query("Dataset == @d").copy()
            mTable = nTable.groupby(["FSel"]).mean()

        spfList = mTable[["AUC_score", "FTime"]]
        R, pval = pearsonr(*zip (*spfList.values))
        R2 = R*R
        print (R, pval)

        fSels = [z[0] for z in spfList.index]
        dSets = [z[1] for z in spfList.index]

        y, x = zip(*spfList.values)
        p, cov = np.polyfit(x, y, 1, cov=True)                     # parameters and covariance from of the fit of 1-D polynom.
        y_model = equation(p, x)                                   # model using the fit parameters; NOTE: parameters here are coefficients

        # Statistics
        n =  len(x)                                          # number of observations
        ps = p.size                                                 # number of parameters
        dof = n - ps                                                # degrees of freedom
        t = stats.t.ppf(0.975, n - ps)                              # used for CI and PI bands

        # Estimates of Error in Data/Model
        resid = y - y_model
        chi2 = np.sum((resid / y_model)**2)                        # chi-squared; estimates error in data
        chi2_red = chi2 / dof                                      # reduced chi-squared; measures goodness of fit
        s_err = np.sqrt(np.sum(resid**2) / dof)                    # standard deviation of the error


        # plot
        if 1 == 1:
            fig, ax = plt.subplots(figsize = (10,10), dpi = DPI)
            sns.scatterplot (x = x, y = y, ax = ax, legend = False)
            ax.plot(x, y_model, "-", color="0.1", linewidth=1.5, alpha=1.0, label="Fit")

            x2 = np.linspace(np.min(x), np.max(x), 100)
            y2 = equation(p, x2)

            plot_ci_manual(t, s_err, n, x, x2, y2, ax=ax)

            pi = t * s_err * np.sqrt(1 + 1/n + (x2 - np.mean(x))**2 / np.sum((x - np.mean(x))**2))
            ax.fill_between(x2, y2 + pi, y2 - pi, color="None", linestyle="--")

            ax.spines["top"].set_color("0.5")
            ax.spines["bottom"].set_color("0.5")
            ax.spines["left"].set_color("0.5")
            ax.spines["right"].set_color("0.5")
            ax.get_xaxis().set_tick_params(direction="out")
            ax.get_yaxis().set_tick_params(direction="out")
            ax.xaxis.tick_bottom()
            ax.yaxis.tick_left()

            plt.xticks(fontsize=20)
            plt.yticks(fontsize=20)
            plt.ylabel('Mean Relative AUC', fontsize = 22, labelpad = 12)
            if applyLOG == True:
                plt.xlabel('Log(1+Mean Training Time)', fontsize= 22, labelpad = 12)
            else:
                plt.xlabel('Mean Training Time', fontsize= 22, labelpad = 12)
            plt.xlim(np.min(x)-0.05, np.max(x) + 0.05)

            right = 0.95
            ypos = 0.93
            legtext = ''
            if len(legtext ) > 0:
                ypos = 0.07
                legtext=legtext+"\n"

            plt.rcParams.update({
                "text.usetex": True,
                "font.family": "sans-serif",
                "font.sans-serif": ["Helvetica"]})
            legpost = ''
            bbox_props = dict(fc="w", ec="0.5", alpha=0.9)
            pTxt = (' = {0:0.2f} ($p$ = {1:0.3f})').format(R2, pval)
            plt.text (right, ypos,
                      (legtext +  "$R^2$" + pTxt),
                      horizontalalignment='right',
                      size = 24, bbox  = bbox_props,
                      transform = ax.transAxes)
            plt.rcParams.update({
                "text.usetex": False,
                "font.family": "sans-serif",
                "font.sans-serif": ["Helvetica"]})

            print ("Bias for", d)
            fig.tight_layout()
            fig.savefig("./results/Training_Times_vs_AUC_"+d+".png", facecolor = 'w')
            if d == "all":
                fig.savefig("./paper/Figure_1C.png", facecolor = 'w')

    plt.close('all')
    pass




if __name__ == "__main__":
    print ("Hi.")

    # load data first
    mlflow.set_tracking_uri(TrackingPath)

    # datasets
    dList = [ "Carvalho2018", "Hosny2018A", "Hosny2018B", "Hosny2018C", "Ramella2018",   "Toivonen2019",
        "Keek2020", "Li2020", "Park2020", "Song2020" , ]

    # obtain results
    results = getResults (dList)

    # timing
    doc = plot_Timings(results, dList, closeDoc = False)
    doc = plot_Timing_Combinations(results, dList, document = doc, closeDoc = True)
    doc = plot_Timing_AUC (results, dList)



#
