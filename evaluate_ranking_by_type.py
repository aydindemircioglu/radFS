#!/usr/bin/python3

from PIL import Image
from PIL import ImageDraw, ImageFont
import json
import seaborn as sns
from matplotlib import cm
from matplotlib.colors import ListedColormap, LinearSegmentedColormap
import matplotlib.colors as colors
from scipy.stats import spearmanr
import pylab
import scipy.cluster.hierarchy as sch


from scipy.stats import pearsonr, friedmanchisquare
from scikit_posthocs import posthoc_nemenyi_friedman
from mlflow.tracking.client import MlflowClient
from mlflow.entities import ViewType


from docx import Document
from docx.shared import Inches

from datetime import datetime
import pandas as pd
import random
import numpy as np
import time
#import pycm
import shutil
import pathlib
import os
import math
import sys
import random
from matplotlib import pyplot
import matplotlib.pyplot as plt
import time
import copy
import random
import pickle
import tempfile
import itertools
import multiprocessing
import socket
from glob import glob
from collections import OrderedDict
import logging
import mlflow
from typing import Dict, Any
import hashlib
import json

from pprint import pprint
from sklearn.metrics import confusion_matrix
from sklearn.metrics import roc_curve, auc, roc_auc_score, precision_recall_curve
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, classification_report

from loadData import *
from utils import *
from parameters import *
from evaluate_utils import *



### parameters
TrackingPath = "/data/results/radFS/mlrun.benchmark"
nCV = 10




def plotDendro (document, zMat, cMat, d = None, Z2 = None, fType = "All"):
    print ('### Plotting Dendro', d, fType)
    plt.rcParams.update({
        "text.usetex": False,
        "font.family": "sans-serif",
        "font.sans-serif": ["Helvetica"]})

    D = zMat.copy()

    if Z2 is None:
        print ("### Obtained a Z2!")
        Y = sch.linkage(cMat, method='single')
        Z2 = sch.dendrogram(Y)
    idx1 = Z2['leaves']
    D = D.iloc[idx1]
    D = D[D.keys()[idx1]]
    fig = pylab.figure(figsize=(13.5,13.5), dpi = 120)

    # Plot distance matrix.
    axmatrix = fig.add_axes([0.3,0.1,0.6,0.6])
    im = axmatrix.matshow(D, norm=MidpointNormalize(midpoint=0,vmin=-1.0, vmax=1.0),\
                aspect='auto', origin='lower', cmap=pylab.cm.RdBu)
    axmatrix.set_yticks([])
    axmatrix.xaxis.tick_bottom()
    xaxis = np.arange(len(D.keys()))
    axmatrix.set_xticks(xaxis)
    axmatrix.set_yticks(xaxis)
    axmatrix.set_xticklabels(D.keys(), rotation = 90,   fontsize = 22)
    axmatrix.set_yticklabels(D.keys(), ha = "right", fontsize = 22)
    plt.rcParams['axes.titlepad'] = 20
    if d == "All":
        axmatrix.set_title ("Ranking for " + fType + " features", fontdict = {'fontsize' : 34})
    else:
        axmatrix.set_title ("Ranking for " + fType + " features on " + d, fontdict = {'fontsize' : 34})
#    plt.tight_layout()


    # Plot colorbar.
    axcolor = fig.add_axes([0.91,0.1,0.02,0.6])
    cbar = pylab.colorbar(im, cax=axcolor)
    cbar.ax.tick_params(labelsize=22)
    #fig.show()

    if d == "All":
        paragraph = document.add_paragraph('All Datasets')
        fig.savefig('results/Ranking_' + fType + '.png', bbox_inches='tight')
        print ("## Saving for ALL!")
        fig.savefig("./results/Figure_5_" + fType + ".png", facecolor = 'w', bbox_inches='tight')
        document.add_picture("./results/Ranking_" + fType + ".png", width=Inches(6.0))
    else:
        paragraph = document.add_paragraph('Dataset ' + d)
        fig.savefig('results/Ranking_' + d + "_" + fType + '.png', bbox_inches='tight')
        document.add_picture("./results/Ranking_"+d+ "_" + fType + ".png", width=Inches(6.0))
    document.add_page_break()

    # for hosny..
    if "Hosny" in d and fType == "All":
        pickle.dump (cMat, open("./results/cmat_" + d + "_" + fType + ".npy","wb"))
        pickle.dump (zMat, open("./results/zmat_" + d + "_" + fType + ".npy","wb"))
        axmatrix.set_title ("Ranking on " + d + " (CV)", fontdict = {'fontsize' : 34})
        fig.savefig('results/Ranking_' + d + "_" + fType + '_CV.png', bbox_inches='tight')

    return Z2



def getRanking (d, fselA, nList):
    fv = []
    for N in sorted(nList, reverse  = True):
        subdata = results.query("Dataset == @d and FSel == @fselA and nFeatures == @N and Clf == 'Constant'")
        z = []
        for m in range(nCV):
            apath = os.path.join(subdata.iloc[0]["Path"], "FPattern_" + str(m) + ".json")
            with open(apath) as f:
                expData = json.load(f)
            z.append(list(expData.values()))
        z = np.asarray(z)
        fv.append(z)
    fv = np.concatenate(fv)

    # we know from stability that the feature names have the correct order
    apath = os.path.join(subdata.iloc[0]["Path"], "FNames_0" + ".json")
    with open(apath) as f:
        fNames = json.load(f)
    fNames = list(fNames.keys())

    featDict = eval (d+"().getFeatureTypes('./data/')")
    fvDict = {}
    fvDict["All"] = np.nanmean(fv, axis = 0)
    for fType in featDict:
        fTypePattern = [u for u, n in enumerate(fNames)  if n in featDict[fType]]
        fvDict[fType] = np.nanmean(fv[:, fTypePattern], axis = 0)
    return fvDict



def plot_Ranking_Matrix (dList, results):
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'
    document.add_heading('Supplemental 6')
    document.add_paragraph(' ')
    document.add_heading('Ranking by feature class', level = 2)
    document.add_paragraph(' ')

    fSels = sorted(list(set(results["FSel"].values)))
    nList = sorted(list(set(results["nFeatures"].values)))

    Z2 = None
    for fType in ["All", "Shape", "Histogram", "Texture" ]:
        fSelList = fSels
        cMatList = {}
        #d = dList[5]
        for d in dList:
            if os.path.exists("./results/rank_" + d + "_" + fType + ".npy") == False:
                rMat = np.ones( (len(fSelList), len(fSelList) ) )
                rMat = pd.DataFrame(rMat, index = fSelList, columns = fSelList)
                preRanks = {}
                for fselA in fSelList:
                    fvA = getRanking(d, fselA, nList)
                    preRanks[fselA] = fvA
                for fselA in fSelList:
                    if fselA == "None":
                        continue
                    for fselB in fSelList:
                        if fselB == "None":
                            continue
                        if fselA == fselB:
                            continue
                        fvA = preRanks[fselA]
                        fvB = preRanks[fselB]
                        if len(set(fvA[fType])) == 1  and len(set(fvB[fType])) == 1:
                            # both constant
                            continue
                        if len(set(fvA[fType])) == 1  or len(set(fvB[fType])) == 1:
                            # now pearson is not defined.
                            # but either we lost all features or we introduced ones even though there we none before
                            # so we count this as completely uncorrelated, 0.0
                            coef, p = 0, 0
                            rMat.at[fselA, fselB] = 0
                        else:
                            rMat.at[fselA, fselB] = spearmanr(fvA[fType], fvB[fType])[0]

                cMat  = rMat.drop(["None"], axis = 1)
                cMat  = cMat.drop(["None"], axis = 0)
                pickle.dump (cMat, open("./results/rank_" + d + "_" + fType + ".npy","wb"))
            else:
                print ("Restoring rank results")
                cMat = pickle.load(open("./results/rank_" + d + "_" + fType + ".npy", "rb"))
            cMatList[d] = cMat

        fullMat = np.nanmean(np.dstack(list(cMatList.values())), axis = 2)
        zMat = cMat.copy()
        zMat[:] = fullMat

        # dendros
        if Z2 is None:
            # make this happens really when we want it
            if  fType != "All":
                print ( fType)
                raise Exception ("Dendro plotting error.")
            Z2 = plotDendro (document, zMat, cMat, d = "All", fType = fType)
            idx1 = Z2['leaves']
            Zkeys = zMat.keys()[idx1]
            pickle.dump (Zkeys, open("./results/Zkeys_ranking.npy", "wb"))
            pickle.dump (Z2, open("./results/Z2_ranking.npy", "wb"))
        for d in dList:
            plotDendro (document, cMatList[d], cMat, d, Z2 = Z2, fType = fType)
        plotDendro (document, zMat, cMat, d = "All", Z2 = Z2, fType = fType)
        plt.close('all')
    document.save('./paper/Supplemental_6.docx')
    plt.close('all')
    pass






def addBorder (img, pos, thickness):
    if pos == "H":
        img = np.hstack([255*np.ones(( img.shape[0],int(img.shape[1]*thickness), 3), dtype = np.uint8),img])
    if pos == "V":
        img = np.vstack([255*np.ones(( int(img.shape[0]*thickness), img.shape[1], 3), dtype = np.uint8),img])
    return img



def addText (finalImage, text = '', org = (0,0), fontFace = '', fontSize = 12, color = (255,255,255)):
     # Convert the image to RGB (OpenCV uses BGR)
     tmpImg = finalImage
     pil_im = Image.fromarray(tmpImg)
     draw = ImageDraw.Draw(pil_im)
     font = ImageFont.truetype(fontFace + ".ttf", fontSize)
     draw.text(org, text, font=font, fill = color)
     tmpImg = np.array(pil_im)
     return (tmpImg.copy())



def join_Ranking_Matrix ():
    import cv2
    fontFace = "Arial"

    imHA = cv2.imread("./results/Figure_5_All.png")
    imHA = addText (imHA, "(a)", (0,0), fontFace, 144, color=(0,0,0))
    imHB = cv2.imread("./results/Figure_5_Histogram.png")
    imHB = addText (imHB, "(c)", (0,0), fontFace, 144, color=(0,0,0))
    imHB = addBorder (imHB, "V", 0.05)
    imgRight = np.vstack([imHA, imHB])
    imgRight = addBorder (imgRight, "H", 0.05)

    imHA = cv2.imread("./results/Figure_5_Shape.png")
    imHA = addText (imHA, "(b)", (0,0), fontFace, 144, color=(0,0,0))
    imHB = cv2.imread("./results/Figure_5_Texture.png")
    imHB = addText (imHB, "(d)", (0,0), fontFace, 144, color=(0,0,0))
    imHB = addBorder (imHB, "V", 0.05)
    imgLeft = np.vstack([imHA, imHB])
    imgLeft = cv2.resize(imgLeft, dsize = imgRight.shape[:2][::-1])
    img = np.hstack([imgRight, imgLeft])
    cv2.imwrite("./paper/Figure_5.png", img)


if __name__ == "__main__":
    print ("Hi.")

    # load data first
    mlflow.set_tracking_uri(TrackingPath)

    # datasets
    dList = [ "Carvalho2018", "Hosny2018A", "Hosny2018B", "Hosny2018C", "Ramella2018",   "Toivonen2019",
        "Keek2020", "Li2020", "Park2020", "Song2020" , ]

    # obtain results
    results = getResults (dList)

     # ranking
    plot_Ranking_Matrix (dList, results)
    join_Ranking_Matrix ()


#
