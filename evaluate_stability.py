#!/usr/bin/python3

import json
import seaborn as sns
from matplotlib import cm
from matplotlib.colors import ListedColormap, LinearSegmentedColormap
import matplotlib.colors as colors
from scipy.stats import spearmanr
import pylab
import scipy.cluster.hierarchy as sch


from scipy.stats import pearsonr, friedmanchisquare
from scikit_posthocs import posthoc_nemenyi_friedman
from mlflow.tracking.client import MlflowClient
from mlflow.entities import ViewType


from docx import Document
from docx.shared import Inches

from datetime import datetime
import pandas as pd
import random
import numpy as np
import time
#import pycm
import shutil
import pathlib
import os
import math
import sys
import random
from matplotlib import pyplot
import matplotlib.pyplot as plt
import time
import copy
import random
import pickle
import tempfile
import itertools
import multiprocessing
import socket
from glob import glob
from collections import OrderedDict
import logging
import mlflow
from typing import Dict, Any
import hashlib
import json

from pprint import pprint
from sklearn.metrics import confusion_matrix
from sklearn.metrics import roc_curve, auc, roc_auc_score, precision_recall_curve
from sklearn.metrics import accuracy_score, f1_score, precision_score, recall_score, classification_report

from loadData import *
from utils import *
from parameters import *
from evaluate_utils import *



### parameters
TrackingPath = "/data/results/radFS/mlrun.benchmark"
nCV = 10





def get_Stability_Table (results, dList):
    fSels = sorted(list(set(results["FSel"].values)))
    nList = sorted(list(set(results["nFeatures"].values)))

    if os.path.exists("./results/stability.feather") == True:
        print ("Restoring stability")
        sTable = pickle.load(open("./results/stability.feather", "rb"))
        return sTable

    print ("Computing stability.")
    sTable = []
    for d in dList:
        data =  eval (d+"().getData('./data/')")
        X = data.drop(["Target"], axis = 1)
        for fsel in fSels:
            if fsel == "None":
                continue

            # list of characteristic vectors
            for N in nList:
                subdata = results.query("Dataset == @d and FSel == @fsel and nFeatures == @N")
                subdata = subdata.sort_values("AUC_mean", ascending = False).iloc[0:1].reset_index(drop = True).copy()
                for idx in range(len(subdata)):
                    row = subdata.iloc[idx]
                    z = []

                    # load the names, make sure it fits
                    apath = os.path.join(row["Path"], "FNames_0" + ".json")
                    with open(apath) as f:
                        fNames = json.load(f)
                    assert (list(X.keys()) == list(fNames.keys()))
                    fNames = list(fNames.keys())

                    for m in range(nCV):
                        apath = os.path.join(row["Path"], "FPattern_" + str(m) + ".json")
                        with open(apath) as f:
                            expData = json.load(f)
                        z.append(list(expData.values()))
                    z = np.asarray(z)
                    m = len(z)
                    SC = {"All": 0}

                    # stability for all
                    for i in range(m):
                        for j in range(i+1,m):
                            coef, p = pearsonr(z[i,:], z[j,:])
                            SC["All"] = SC["All"] + coef
                    SC["All"] = 2/(m*(m-1))*SC["All"]


                    featDict = eval (d+"().getFeatureTypes('./data/')")
                    for fType in featDict:
                        SC[fType] = 0
                        fTypePattern = [u for u, n in enumerate(fNames)  if n in featDict[fType]]
                        if len(fTypePattern) == 0:
                            SC[fType] = np.nan
                            continue
                        nPairs = 0
                        for i in range(m):
                            for j in range(i+1,m):
                                # check if we have any selection at all
                                if len(set(z[i,fTypePattern])) == 1  and len(set(z[j,fTypePattern])) == 1:
                                    continue
                                nPairs = nPairs + 1
                                if len(set(z[i,fTypePattern])) == 1  or len(set(z[j,fTypePattern])) == 1:
                                    # now pearson is not defined.
                                    # but either we lost all features or we introduced ones even though there we none before
                                    # so we count this as completely uncorrelated, 0.0
                                    coef, p = 0, 0
                                else:
                                    coef, p = pearsonr(z[i,fTypePattern], z[j,fTypePattern])
                                SC[fType] = SC[fType] + coef
                        # just check
                        if nPairs == 0:
                            SC[fType] = np.nan
                        else:
                            SC[fType] = 2/(m*(m-1))*SC[fType]

                    sRow = {"FeatureSelection": fsel, "Dataset": d, "NFeatures": N, "Stability": SC["All"], "AUC": row["AUC_mean"]}
                    for fType in featDict:
                        sRow["SC_"+fType] = SC[fType]

                    sTable.append(sRow)

    sTable = pd.DataFrame(sTable)
    print ("Pickling stability.")
    pickle.dump (sTable, open("./results/stability.feather","wb"))
    return sTable



def plot_Stability(results, dList, sTable):
    # save to results ..
    sOrder = None
    for fType in ["All"]:
        document = Document()
        font = document.styles['Normal'].font
        font.name = 'Arial'
        if fType == "All":
            document.add_heading('Supplemental 3')
            sKey = "Stability"
            yaxLabel = "Stability"
        else:
            document.add_heading('Supplemental 7')
            document.add_paragraph(' ')
            document.add_heading('Stability by feature class', level = 2)
            document.add_paragraph(' ')
            sKey = "SC_" + str(fType)
            yaxLabel = "Stability [" + str(fType) + "]"
        document.add_paragraph(' ')
        paragraph = document.add_paragraph(fType + "Features")
        for d in dList + ["all"]:
            paragraph = document.add_paragraph('Dataset ' + d)
            if d == "all":
                dTable = sTable.copy()
            else:
                dTable = sTable[sTable["Dataset"] == d].copy()

            sOrder = dTable.groupby (["FeatureSelection"])[sKey].median()
            sOrder = sOrder.sort_values(ascending = False)

            # create plot
            plotData = dTable.copy()
            fig, ax = plt.subplots(figsize = (20,10), dpi = DPI)
            sns.set(style='white')
            sns.boxplot(x = 'FeatureSelection', y = sKey, data= plotData,   order = sOrder.keys())
            print ("PLOT WITH", plotData.shape)
            sns.despine()

            missings = np.isfinite(sTable["SC_Histogram"])
            missings = np.sum(missings)
            missings/sTable.shape[0]
            missings = np.isfinite(sTable["SC_Texture"])
            missings = np.sum(missings)
            missings/sTable.shape[0]

            ax.set_xticklabels(sOrder.keys(), rotation = 45, ha = "right", fontsize = 22)
            #ax.set_yscale('log')
            ax.yaxis.set_tick_params ( labelsize= 22)
            ax.set_xlabel ("Feature Selection Method", fontsize = 26)
            ax.set_ylabel (yaxLabel + "(Pearson Correlation)", fontsize = 26)
            #ax.set_title("Feature Selection Stability", fontsize = 24)
            plt.tight_layout()
            fig.savefig("./results/Stability_"+d+"_"+ fType + ".png", facecolor = 'w')
            if d == "all":
                if fType == "All":
                    fig.savefig("./paper/Figure_3.png", facecolor = 'w')
                else:
                    fig.savefig("./results/Figure_3_" + fType + ".png", facecolor = 'w')
            document.add_picture("./results/Stability_"+d+"_"+ fType + ".png", width=Inches(6.0))
            document.add_page_break()

        document.save('./paper/Supplemental_3.docx')
        plt.close('all')
    pass




def plot_Stability_by_Type(results, dList, sTable):
    document = Document()
    font = document.styles['Normal'].font
    font.name = 'Arial'
    document.add_heading('Supplemental 4')
    document.add_paragraph(' ')
    document.add_heading('Stability by feature class', level = 2)
    document.add_paragraph(' ')

    # save to results ..
    sOrder = sTable.groupby (["FeatureSelection"])["Stability"].median()
    sOrder = sOrder.sort_values(ascending = False)

    fTypes = ["All", "Histogram", "Shape", "Texture"]
    rMatList = {}
    for d in dList:
        rMat = np.zeros( (len(sOrder.index), len(fTypes) ) )
        rMat = pd.DataFrame(rMat, index = sOrder.index, columns = fTypes)
        for fType in fTypes:
            dTable = sTable[sTable["Dataset"] == d].copy()
            if fType == "All":
                sKey = "Stability"
            else:
                sKey = "SC_" + fType
            vals = dTable.groupby(["FeatureSelection"])[sKey].median()
            for v in range(len(vals)):
                rMat.at[vals.index[v], fType] = vals[v]
        rMatList[d] = rMat

    for d in ["All"] + dList:
        paragraph = document.add_paragraph('Dataset ' + d)
        if d == "All":
            rMat = np.stack([rMatList[r] for r in rMatList])
            rMat = np.nanmean(rMat, axis = 0)
            cMat = rMatList["Song2020"].copy()
            rMat = pd.DataFrame(rMat, index = cMat.index, columns = cMat.keys())
        else:
            rMat = rMatList[d]
            cMat = rMatList[d].copy()
            rMat = pd.DataFrame(rMat, index = cMat.index, columns = cMat.keys())
        labels = rMat.copy()
        labels = np.array(labels)
        labels = labels.round(2).astype(str)

        if 1 == 1:
            fig, ax = plt.subplots(figsize = (10,15), dpi = DPI)
            sns.set(style='white')
            sns.heatmap(rMat, annot = labels,  cmap = "Reds", fmt = '', annot_kws={"fontsize":21}, cbar = False)
            ax.set_xticklabels(rMat.keys(), rotation = 45, ha = "right", fontsize = 21)
            ax.set_yticklabels(rMat.index, rotation = 0, ha = "right", fontsize = 21)
            ax.yaxis.set_tick_params ( labelsize= 21)
            ax.set_xlabel ("", fontsize = 19)
            ax.set_ylabel ("", fontsize = 19)
            ax.set_title("", fontsize = 24)
            plt.tight_layout()
            if d == "All":
                fig.savefig("./paper/Figure_4.png", facecolor = 'w')
            fig.savefig("./results/Stability_by_Type_"+d+".png", facecolor = 'w')
            document.add_picture("./results/Stability_by_Type_"+d+".png", width=Inches(4.5))
            document.add_page_break()

    document.save('./paper/Supplemental_4.docx')
    plt.close('all')
    return document

    pass





def plot_Stability_AUC(results, dList, sTable):
    nTable = []
    for d in dList:
        subdata = sTable.query("Dataset == @d")
        subdata = subdata.sort_values("AUC", ascending = False).iloc[0:1].reset_index(drop = True).copy()
        lmax = subdata.iloc[0]["AUC"]
        subdata = sTable.query("Dataset == @d").copy()
        subdata["AUC_score"] = lmax - subdata["AUC"]
        nTable.append(subdata)
    mainTable = pd.concat(nTable)

    for d in dList + ["all"]:
        if d == "all":
            mTable = mainTable.groupby(["FeatureSelection", "Dataset"]).mean()
        else:
            nTable = mainTable.query("Dataset == @d").copy()
            mTable = nTable.groupby(["FeatureSelection"]).mean()

        spfList = mTable[["AUC_score", "Stability"]]
        R, pval = pearsonr(*zip (*spfList.values))
        R2 = R*R
        print (R, pval)

        fSels = [z[0] for z in spfList.index]
        dSets = [z[1] for z in spfList.index]

        y, x = zip(*spfList.values)
        p, cov = np.polyfit(x, y, 1, cov=True)                     # parameters and covariance from of the fit of 1-D polynom.
        y_model = equation(p, x)                                   # model using the fit parameters; NOTE: parameters here are coefficients

        # Statistics
        n =  len(x)                                          # number of observations
        ps = p.size                                                 # number of parameters
        dof = n - ps                                                # degrees of freedom
        t = stats.t.ppf(0.975, n - ps)                              # used for CI and PI bands

        # Estimates of Error in Data/Model
        resid = y - y_model
        chi2 = np.sum((resid / y_model)**2)                        # chi-squared; estimates error in data
        chi2_red = chi2 / dof                                      # reduced chi-squared; measures goodness of fit
        s_err = np.sqrt(np.sum(resid**2) / dof)                    # standard deviation of the error


        # plot
        if 1 == 1:
            fig, ax = plt.subplots(figsize = (10,10), dpi = DPI)
            sns.scatterplot (x = x,y = y,  ax = ax)
            ax.plot(x, y_model, "-", color="0.1", linewidth=1.5, alpha=1.0, label="Fit")

            x2 = np.linspace(np.min(x), np.max(x), 100)
            y2 = equation(p, x2)

            # Confidence Interval (select one)
            plot_ci_manual(t, s_err, n, x, x2, y2, ax=ax)

            # Prediction Interval
            pi = t * s_err * np.sqrt(1 + 1/n + (x2 - np.mean(x))**2 / np.sum((x - np.mean(x))**2))
            ax.fill_between(x2, y2 + pi, y2 - pi, color="None", linestyle="--")


            # Figure Modifications --------------------------------------------------------
            # Borders
            ax.spines["top"].set_color("0.5")
            ax.spines["bottom"].set_color("0.5")
            ax.spines["left"].set_color("0.5")
            ax.spines["right"].set_color("0.5")
            ax.get_xaxis().set_tick_params(direction="out")
            ax.get_yaxis().set_tick_params(direction="out")
            ax.xaxis.tick_bottom()
            ax.yaxis.tick_left()
            plt.xticks(fontsize=20)
            plt.yticks(fontsize=20)
            plt.ylabel('Mean Relative AUC', fontsize = 22, labelpad = 12)
            plt.xlabel('Mean Stability', fontsize= 22, labelpad = 12)
            plt.xlim(np.min(x)-0.05, np.max(x) + 0.05)
            #plt.ylim(np.min(y), 1.025)

            right = 0.95
            ypos = 0.93
            legtext = ''
            if len(legtext ) > 0:
                ypos = 0.07
                legtext=legtext+"\n"

            plt.rcParams.update({
                "text.usetex": True,
                "font.family": "sans-serif",
                "font.sans-serif": ["Helvetica"]})
            legpost = ''
            bbox_props = dict(fc="w", ec="0.5", alpha=0.9)
            pTxt = (' = {0:0.2f} ($p$ = {1:0.3f})').format(R2, pval)
            plt.text (right, ypos,
                      (legtext +  "$R^2$" + pTxt),
                      horizontalalignment='right',
                      size = 24, bbox  = bbox_props,
                      transform = ax.transAxes)
            plt.rcParams.update({
                "text.usetex": False,
                "font.family": "sans-serif",
                "font.sans-serif": ["Helvetica"]})

            print ("Bias for", d)
            fig.tight_layout()
            fig.savefig("./results/Stability_vs_AUC_" + d + ".png", facecolor = 'w')
            if d == "all":
                fig.savefig("./paper/Figure_2B.png", facecolor = 'w')

    plt.close('all')
    pass







if __name__ == "__main__":
    print ("Hi.")

    # load data first
    mlflow.set_tracking_uri(TrackingPath)

    # datasets
    dList = [ "Carvalho2018", "Hosny2018A", "Hosny2018B", "Hosny2018C", "Ramella2018",   "Toivonen2019",
        "Keek2020", "Li2020", "Park2020", "Song2020" , ]

    # obtain results
    results = getResults (dList)

    # stability
    sTable = get_Stability_Table (results, dList)

    plot_Stability_by_Type(results, dList, sTable)
    plot_Stability(results, dList, sTable)
    plot_Stability_AUC (results, dList, sTable)


#
